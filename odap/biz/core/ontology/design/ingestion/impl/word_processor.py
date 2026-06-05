import logging
from typing import Any, Dict, List, Optional, Union

logger = logging.getLogger(__name__)

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class WordProcessor:
    def extract_text(self, file_path_or_bytes: Union[str, bytes]) -> Dict[str, Any]:
        if not DOCX_AVAILABLE:
            return self._fallback_extract_text(file_path_or_bytes)

        try:
            if isinstance(file_path_or_bytes, bytes):
                import io
                source = io.BytesIO(file_path_or_bytes)
            else:
                source = file_path_or_bytes

            doc = docx.Document(source)
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            full_text = "\n".join(paragraphs)
            return {
                "status": "success",
                "text": full_text,
                "paragraph_count": len(paragraphs),
                "source_type": "word",
            }
        except Exception as e:
            logger.warning("Word text extraction failed: %s", e)
            return self._fallback_extract_text(file_path_or_bytes)

    def extract_tables(self, file_path_or_bytes: Union[str, bytes]) -> Dict[str, Any]:
        if not DOCX_AVAILABLE:
            return {"status": "fallback", "tables": [], "message": "python-docx not available"}

        try:
            if isinstance(file_path_or_bytes, bytes):
                import io
                source = io.BytesIO(file_path_or_bytes)
            else:
                source = file_path_or_bytes

            doc = docx.Document(source)
            all_tables = []

            for table_idx, table in enumerate(doc.tables):
                rows_data = []
                for row in table.rows:
                    row_cells = [cell.text for cell in row.cells]
                    rows_data.append(row_cells)

                if rows_data:
                    headers = rows_data[0]
                    rows = rows_data[1:]
                    all_tables.append({
                        "table_index": table_idx + 1,
                        "headers": headers,
                        "rows": rows,
                        "row_count": len(rows),
                    })

            return {
                "status": "success",
                "tables": all_tables,
                "table_count": len(all_tables),
                "source_type": "word",
            }
        except Exception as e:
            logger.warning("Word table extraction failed: %s", e)
            return {"status": "error", "tables": [], "message": str(e)}

    def _fallback_extract_text(self, file_path_or_bytes: Union[str, bytes]) -> Dict[str, Any]:
        if isinstance(file_path_or_bytes, bytes):
            try:
                text = file_path_or_bytes.decode("utf-8", errors="replace")
                return {
                    "status": "fallback",
                    "text": text,
                    "paragraph_count": 1,
                    "source_type": "word",
                    "message": "python-docx not available, raw bytes decoded",
                }
            except Exception:
                return {"status": "error", "text": "", "paragraph_count": 0, "source_type": "word", "message": "python-docx not available"}

        try:
            with open(file_path_or_bytes, "rb") as f:
                raw = f.read()
            text = raw.decode("utf-8", errors="replace")
            return {
                "status": "fallback",
                "text": text,
                "paragraph_count": 1,
                "source_type": "word",
                "message": "python-docx not available, raw file decoded",
            }
        except Exception as e:
            return {"status": "error", "text": "", "paragraph_count": 0, "source_type": "word", "message": str(e)}
